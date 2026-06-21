"""
结构化文档解析器（适配器模式）

所有格式解析器输出统一的Block列表
"""
import os
import sys
import logging
from typing import List, Optional
from abc import ABC, abstractmethod

from knowledge.block import Block
from knowledge.document_cleaner import DocumentCleaner

logger = logging.getLogger(__name__)


class DocumentParser(ABC):
    """文档解析器基类（适配器接口）"""

    @abstractmethod
    def parse(self, file_path: str) -> List[Block]:
        """
        解析文档，返回Block列表

        Args:
            file_path: 文件路径

        Returns:
            Block列表
        """
        pass


class WordParser(DocumentParser):
    """Word文档解析器（.docx）"""

    def parse(self, file_path: str) -> List[Block]:
        """解析Word文档"""
        try:
            from docx import Document
            from docx.enum.text import WD_STYLE_TYPE
        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")

        doc = Document(file_path)
        blocks = []
        current_section = []
        filename = os.path.basename(file_path)

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # 判断块类型
            style_name = para.style.name.lower() if para.style else ""

            if 'title' in style_name:
                block_type = 'title'
            elif 'heading' in style_name:
                block_type = 'heading'
                # 提取标题层级
                level = 1
                if 'heading' in style_name:
                    parts = style_name.split()
                    for part in parts:
                        if part.isdigit():
                            level = int(part)
                            break
                current_section.append((level, text))
            else:
                block_type = 'paragraph'

            # 构建章节路径
            section_path = ' > '.join([sec[1] for sec in current_section])

            blocks.append(Block(
                type=block_type,
                text=text,
                page=None,  # Word没有页码概念
                source_file=filename,
                section_path=section_path if section_path else None,
                metadata={'line_number': idx}
            ))

        # 提取表格
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(' | '.join(cells))

            table_text = '\n'.join(rows)
            if table_text.strip():
                blocks.append(Block(
                    type='table',
                    text=table_text,
                    page=None,
                    source_file=filename,
                    section_path=' > '.join([sec[1] for sec in current_section]),
                    metadata={'table_index': table_idx, 'row_count': len(rows)}
                ))

        logger.info(f"Word解析完成: {len(blocks)} 个块")
        return blocks


class MarkdownParser(DocumentParser):
    """Markdown文档解析器（.md）"""

    def parse(self, file_path: str) -> List[Block]:
        """解析Markdown文档"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        blocks = []
        current_section = []
        filename = os.path.basename(file_path)
        lines = content.split('\n')

        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            # 标题
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                title_text = line.lstrip('#').strip()

                # 更新章节路径
                while current_section and current_section[-1][0] >= level:
                    current_section.pop()
                current_section.append((level, title_text))

                blocks.append(Block(
                    type='heading' if level > 1 else 'title',
                    text=title_text,
                    page=None,
                    source_file=filename,
                    section_path=' > '.join([sec[1] for sec in current_section]),
                    metadata={'heading_level': level, 'line_number': i}
                ))
                i += 1

            # 表格（简单检测：包含 |）
            elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
                table_lines = [line]
                i += 1
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i])
                    i += 1

                table_text = '\n'.join(table_lines)
                blocks.append(Block(
                    type='table',
                    text=table_text,
                    page=None,
                    source_file=filename,
                    section_path=' > '.join([sec[1] for sec in current_section]),
                    metadata={'line_number': i - len(table_lines)}
                ))

            # 普通段落
            else:
                para_lines = [line]
                i += 1
                while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and '|' not in lines[i]:
                    para_lines.append(lines[i].strip())
                    i += 1

                para_text = ' '.join(para_lines)
                blocks.append(Block(
                    type='paragraph',
                    text=para_text,
                    page=None,
                    source_file=filename,
                    section_path=' > '.join([sec[1] for sec in current_section]),
                    metadata={'line_number': i - len(para_lines)}
                ))

        logger.info(f"Markdown解析完成: {len(blocks)} 个块")
        return blocks


class PDFParser(DocumentParser):
    """PDF文档解析器（文字版）"""

    def parse(self, file_path: str) -> List[Block]:
        """解析文字版PDF"""
        try:
            import fitz  # PyMuPDF
        except ImportError as e:
            error_msg = (
                f"PyMuPDF 未安装。请运行: {sys.executable} -m pip install pymupdf\n"
                "注意：如果启动日志显示 PyMuPDF 可用但此处仍报错，说明当前代码运行的解释器与启动检查的不一致。"
            )
            raise ImportError(error_msg) from e

        doc = fitz.open(file_path)
        filename = os.path.basename(file_path)
        pages_text = []

        # 提取每页文本
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()
            pages_text.append(page_text)

        # 清洗文档（删除页眉/页脚/页码）
        cleaned_text = DocumentCleaner.clean_document(pages_text)

        # 简单分块（按段落）
        blocks = []
        paragraphs = [p.strip() for p in cleaned_text.split('\n\n') if p.strip()]

        for idx, para in enumerate(paragraphs):
            # 简单判断：大写或粗体 → 标题
            block_type = 'heading' if para.isupper() or len(para) < 50 else 'paragraph'

            blocks.append(Block(
                type=block_type,
                text=para,
                page=None,  # TODO: 精确页码需要更复杂的逻辑
                source_file=filename,
                section_path=None,
                metadata={'para_index': idx}
            ))

        logger.info(f"PDF解析完成: {len(blocks)} 个块")
        doc.close()
        return blocks


class StructuredDocumentLoader:
    """结构化文档加载器（门面）"""

    def __init__(self):
        self.parsers = {
            'docx': WordParser(),
            'md': MarkdownParser(),
            'markdown': MarkdownParser(),
            'pdf': PDFParser(),
        }

    def load_document(self, file_path: str, file_type: Optional[str] = None) -> List[Block]:
        """
        加载文档，返回Block列表

        Args:
            file_path: 文件路径
            file_type: 文件类型（可选）

        Returns:
            Block列表
        """
        if file_type is None:
            _, ext = os.path.splitext(file_path)
            file_type = ext.lower().lstrip('.')

        parser = self.parsers.get(file_type)
        if not parser:
            raise ValueError(f"不支持的文件类型: {file_type}")

        try:
            blocks = parser.parse(file_path)
            logger.info(f"文档加载成功: {len(blocks)} 个块")
            return blocks
        except Exception as e:
            logger.exception(f"文档加载失败: {e}")
            raise
