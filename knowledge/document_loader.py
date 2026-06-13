"""
文档加载器

支持多种格式文档的解析和文本提取
"""
import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class DocumentLoader:
    """文档加载器"""

    @staticmethod
    def load_pdf(file_path: str) -> str:
        """
        加载 PDF 文档

        Args:
            file_path: PDF 文件路径

        Returns:
            提取的文本内容
        """
        try:
            import pypdf

            text_parts = []
            with open(file_path, 'rb') as file:
                reader = pypdf.PdfReader(file)
                for page_num, page in enumerate(reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"[第 {page_num} 页]\n{page_text}")

            return "\n\n".join(text_parts)

        except ImportError:
            logger.warning("pypdf 未安装，尝试使用 pdfplumber")
            try:
                import pdfplumber

                text_parts = []
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_parts.append(f"[第 {page_num} 页]\n{page_text}")

                return "\n\n".join(text_parts)

            except ImportError:
                raise ImportError("需要安装 pypdf 或 pdfplumber: pip install pypdf pdfplumber")

    @staticmethod
    def load_docx(file_path: str) -> str:
        """
        加载 Word 文档

        Args:
            file_path: Word 文件路径

        Returns:
            提取的文本内容
        """
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)

            result = "\n\n".join(paragraphs)

            if not result.strip():
                raise ValueError("Word 文档内容为空，可能是加密或损坏的文件")

            return result

        except ImportError:
            raise ValueError("Word 文档解析需要 python-docx 库。请安装: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Word 文档解析失败: {str(e)}。请确保文件未损坏且未加密")

    @staticmethod
    def load_markdown(file_path: str) -> str:
        """
        加载 Markdown 文档

        Args:
            file_path: Markdown 文件路径

        Returns:
            文本内容
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    @staticmethod
    def load_text(file_path: str) -> str:
        """
        加载纯文本文档

        Args:
            file_path: 文本文件路径

        Returns:
            文本内容
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    @staticmethod
    def load_document(file_path: str, file_type: Optional[str] = None) -> str:
        """
        自动识别并加载文档

        Args:
            file_path: 文件路径
            file_type: 文件类型（可选，如 'pdf', 'docx', 'md', 'txt'）

        Returns:
            提取的文本内容

        Raises:
            ValueError: 不支持的文件类型
        """
        if file_type is None:
            _, ext = os.path.splitext(file_path)
            file_type = ext.lower().lstrip('.')

        loaders = {
            'pdf': DocumentLoader.load_pdf,
            'docx': DocumentLoader.load_docx,
            'md': DocumentLoader.load_markdown,
            'markdown': DocumentLoader.load_markdown,
            'txt': DocumentLoader.load_text,
            'text': DocumentLoader.load_text
        }

        loader = loaders.get(file_type)
        if not loader:
            raise ValueError(f"不支持的文件类型: {file_type}，支持的格式: {', '.join(loaders.keys())}")

        try:
            return loader(file_path)
        except Exception as e:
            logger.error(f"加载文档失败: {e}")
            raise
